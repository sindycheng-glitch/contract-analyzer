import os
import io
import json
from typing import List

from google import genai
from google.genai import types
from docx import Document
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title="Contract Analyzer API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["POST", "GET"],
    allow_headers=["*"],
)

# ── Single contract analysis prompt ──────────────────────────────────────────
ANALYSIS_PROMPT = """你是一位嚴格且經驗豐富的企業法務顧問。請仔細審閱以下合約內容，針對下列三個面向進行風險評估：

1. **付款與財務條款** (payment_terms)：付款期限、逾期罰款、價格調整、預付款、發票規範
2. **責任限制與賠償條款** (liability_indemnification)：責任上限、間接損失排除、賠償範圍、保險要求、免責聲明
3. **終止與違約條款** (termination_breach)：終止權利、違約定義、通知期限、終止後義務、競業禁止

對每個面向：
- 評估風險等級：safe（條款合理，對我方保護充分）、caution（有潛在風險，建議談判修改）、critical（高風險條款，強烈建議拒絕或重新談判）
- 引用合約中最相關的原文條款（若無相關條款請標示「合約未包含此條款，建議補充」）
- 給出簡短法務評語與具體修改建議

整體風險以最高的單項風險為準（任一 critical → overall critical；無 critical 但有 caution → overall caution；全部 safe → overall safe）。

請嚴格以下列 JSON 格式回覆，不要包含任何 markdown 標記或額外說明：
{
  "overall_risk": "safe|caution|critical",
  "summary": "整體合約風險摘要，說明主要風險點（2-3句話）",
  "analyses": [
    {
      "id": "payment_terms",
      "name": "付款與財務條款",
      "risk_level": "safe|caution|critical",
      "original_clause": "合約原文摘錄（或說明缺失）",
      "comment": "法務評語（指出風險所在）",
      "suggestion": "具體修改建議（可直接用於與廠商溝通的文字）"
    },
    {
      "id": "liability_indemnification",
      "name": "責任限制與賠償條款",
      "risk_level": "safe|caution|critical",
      "original_clause": "合約原文摘錄（或說明缺失）",
      "comment": "法務評語（指出風險所在）",
      "suggestion": "具體修改建議（可直接用於與廠商溝通的文字）"
    },
    {
      "id": "termination_breach",
      "name": "終止與違約條款",
      "risk_level": "safe|caution|critical",
      "original_clause": "合約原文摘錄（或說明缺失）",
      "comment": "法務評語（指出風險所在）",
      "suggestion": "具體修改建議（可直接用於與廠商溝通的文字）"
    }
  ]
}"""

# ── Multi-contract comparison prompt ─────────────────────────────────────────
COMPARISON_PROMPT = """你是公司的一人法務，有大量的合約需審閱，有時是新約有時是續約新增條款。

請仔細比對上方提供的各份合約，針對以下面向找出所有關鍵差異：
1. 付款與財務條款（付款期限、金額、逾期罰款等）
2. 責任限制與賠償條款（責任上限、賠償範圍、保險等）
3. 終止與違約條款（終止條件、通知期限、違約責任等）
4. 保密條款（保密範圍、期限、例外情況等）
5. 智慧財產權與成果歸屬
6. 其他有重大差異的條款

對每個差異點，請明確指出：
- 哪個版本對甲方（買方／委託方／發包方）較有利，並引用該版本的條款原文
- 哪個版本對乙方（賣方／承接方／廠商）較有利，並引用該版本的條款原文
- 各方評語：一句話說明為何此版本有利（務必具體說明利益所在）
- AI修改建議：建議的平衡條款文字，可直接複製給廠商協商使用

注意：
- 若某合約完全未載明某條款，仍應列出此差異，並在缺失方標示「本合約未載明此條款」
- 每個差異點應聚焦在單一具體的條款差異，不要混合多個議題
- 至少找出 5 個以上的差異點

請嚴格以下列 JSON 格式回覆，不要包含任何 markdown 標記或額外說明：
{
  "summary": "整體比較摘要（3-4句話，說明各份合約的主要差異方向與整體風險差距）",
  "differences": [
    {
      "category": "條款類別（如：付款條款）",
      "description": "此差異點的一句話核心說明",
      "primary_beneficiary": "party_a 或 party_b（此差異整體上對誰更有利）",
      "party_a_favorable": {
        "source": "來源合約的完整檔名",
        "clause": "對甲方有利的條款原文（或「本合約未載明此條款」）",
        "comment": "為何對甲方有利（1-2句，需具體）"
      },
      "party_b_favorable": {
        "source": "來源合約的完整檔名",
        "clause": "對乙方有利的條款原文（或「本合約未載明此條款」）",
        "comment": "為何對乙方有利（1-2句，需具體）"
      },
      "ai_suggestion": "建議的平衡條款文字，可直接用於與廠商協商（完整可執行的條款文字）"
    }
  ]
}"""


def parse_response(text: str) -> dict:
    text = text.strip()
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    return json.loads(text.strip())


async def extract_file_content(file: UploadFile) -> tuple[str, bytes | str]:
    """Returns (type, content) where type is 'pdf' or 'text'."""
    filename = (file.filename or "").lower()
    content = await file.read()

    if not content:
        raise HTTPException(status_code=400, detail=f"檔案 {file.filename} 是空的")

    if filename.endswith(".pdf"):
        return "pdf", content

    if filename.endswith(".docx"):
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
        if not text.strip():
            raise HTTPException(status_code=400, detail=f"無法從 {file.filename} 提取文字")
        return "text", text

    raise HTTPException(
        status_code=400,
        detail=f"不支援的檔案格式：{file.filename}（請使用 .pdf 或 .docx）",
    )


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/analyze")
async def analyze_contract(file: UploadFile = File(...)):
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY 環境變數未設定")

    client = genai.Client(api_key=api_key)
    content_type, content = await extract_file_content(file)

    try:
        if content_type == "pdf":
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[
                    types.Part.from_bytes(data=content, mime_type="application/pdf"),
                    ANALYSIS_PROMPT,
                ],
            )
        else:
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=f"以下是合約全文：\n\n{content}\n\n---\n\n{ANALYSIS_PROMPT}",
            )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 分析失敗：{str(e)}")

    try:
        return parse_response(response.text)
    except (json.JSONDecodeError, KeyError) as e:
        raise HTTPException(status_code=500, detail=f"AI 回應格式錯誤：{str(e)}")


@app.post("/compare")
async def compare_contracts(files: List[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="請至少上傳 2 份合約進行比對")

    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY 環境變數未設定")

    client = genai.Client(api_key=api_key)

    contents: list = ["以下是需要進行比對的合約內容：\n\n"]

    for i, file in enumerate(files):
        content_type, content = await extract_file_content(file)
        contents.append(f"=== 合約 {i + 1}：{file.filename} ===\n")
        if content_type == "pdf":
            contents.append(types.Part.from_bytes(data=content, mime_type="application/pdf"))
        else:
            contents.append(content)
        contents.append("\n\n")

    contents.append(COMPARISON_PROMPT)

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=contents,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 分析失敗：{str(e)}")

    try:
        return parse_response(response.text)
    except (json.JSONDecodeError, KeyError) as e:
        raise HTTPException(status_code=500, detail=f"AI 回應格式錯誤：{str(e)}")
